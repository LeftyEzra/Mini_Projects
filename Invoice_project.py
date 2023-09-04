#from docxtpl import DocxTemplate
from docx import Document
#import jinja2 as jj
from docx.shared import Inches

def client_invoice(name,unit ,description,price):
    document = Document()
    document.add_picture("Tinanzze.PNG",width=Inches(1))
    document.add_heading("Invoice",0)
    document.add_heading("Company", 0)
    p1= document.add_paragraph('Dear ')
    p1.add_run(name).bold=True
    p1.add_run(',')

    p2 = document.add_paragraph("Please find the attached invoice for your recent purchase ")
    p2.add_run(str(f"{unit} ")).bold=True
    p2.add_run('units of ')
    p2.add_run(description).bold=True
    p2.add_run('.')

    [document.add_paragraph("")for _ in range(2)]

    table = document.add_table(rows=1,cols=4)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Quantity'
    hdr_cells[1].text = 'Description'
    hdr_cells[2].text = ' Unit Price'
    hdr_cells[3].text = 'Total'

    for num in range(4):
        hdr_cells[num].paragraphs[0].runs[0].font.bold=True

    row_cells = table.add_row().cells
    row_cells[0].text = f"{unit:.2f}"
    row_cells[1].text = description
    row_cells[2].text = f"{price:.2f}"
    row_cells[3].text = f"{unit * price:,.2f}"

    [document.add_paragraph("") for _ in range(2)]

    table = document.add_table(rows=1, cols=4)
    hdr_cells = table.rows[0].cells
    #hdr_cells[0].text = ''
    #hdr_cells[1].text = ''
    hdr_cells[2].text = ' Subtotal'
    hdr_cells[3].text =  row_cells[3].text

    hdr_cells[2].paragraphs[0].runs[0].font.bold=True


    [document.add_paragraph("") for _ in range(10)]

    document.add_paragraph("We appreciate your patronage")
    document.add_paragraph("Sincerely")
    document.add_paragraph("Lefty")

    document.save(f"{name}.docx")

client_invoice('Lefty Ezra',10,'pen',100)

