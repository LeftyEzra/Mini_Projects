"""import tkinter
from tkinter import ttk
window = tkinter.Tk()

columns = ("Name","Age", "Subscription", "Employment")

style = ttk.Style()
style.configure("mystyle.Treeview", highlightthickness=0, bd=0, font=('Calibri', 11)) # Modify the font of the body
style.configure("mystyle.Treeview.Heading", font=('Calibri', 13,'bold')) # Modify the font of the headings
style.layout("mystyle.Treeview", [('mystyle.Treeview.treearea', {'sticky': 'nswe'})]) # Remove the borders

tree=ttk.Treeview(window,show="headings",columns=columns,height=15)
tree.pack()


window.mainloop()"""
#import os
#import time

"""import qrcode

qr = qrcode.QRCode(version=1,
                   error_correction=qrcode.constants.ERROR_CORRECT_L,
                   box_size=15,
                   border=2,)               # Specify the type of box you want to generate

qr.add_data("https://youtu.be/Ao-MdwTM51Y") #The data,url, image,text, or anything you want to pass as message
qr.make(fit=True)                           #command to create/generate  qrcode

img = qr.make_image(fill_color="red",back_ground="black") # The background and the foreground color of the qrcode
img.save("IMG.PNG")                                       #save  the code as an image"""#Qr code
"""from docx2pdf import convert

convert("my_invoice.docx")
convert("my_invoice.docx", "my.pdf")
from docx import Document
#import paragraph

document = Document()
#document.add_heading("Sales Receipt",0)
#document.add_paragraph()
#document.add_heading("Tinazze-foot",0)
#document.add_picture("Tinanzze.PNG")

document.save("txt.docx")"""#Convert docx to pdf
#path = "C:\\Users\\user\\Desktop\\DATA ANALYTICS.txt"
#if os.path.exists(path):
    #print("This is a same file")

#else:
    #print("That location is not same")

#file = lambda path: "this file exist" if os.path.exists(path) else "Doesnt exit".file(path)
#print(file(path))
